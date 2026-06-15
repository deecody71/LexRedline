"""DOCX contract parser using python-docx."""
import re
from typing import List, Optional
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from src.models import Contract, Section
from src.parsers.base import BaseParser


class DOCXParser(BaseParser):
    """Parse DOCX contracts into structured Contract models."""

    def parse(self, file_path: str) -> Contract:
        """Parse a DOCX file and extract structured content."""
        doc = DocxDocument(file_path)
        return self._extract(doc, file_path)

    def parse_bytes(self, content: bytes, filename: str) -> Contract:
        """Parse DOCX from bytes."""
        import io
        doc = DocxDocument(io.BytesIO(content))
        return self._extract(doc, filename)

    def _extract(self, doc, filename: str) -> Contract:
        """Extract structured content from a python-docx Document."""
        full_text_parts = []
        sections: List[Section] = []
        metadata = {}

        # Extract document properties
        try:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.created:
                metadata["created"] = str(props.created)
            if props.modified:
                metadata["modified"] = str(props.modified)
        except Exception:
            pass

        # Count paragraphs as page approximation (DOCX doesn't store pages)
        paragraph_count = len(doc.paragraphs)

        current_section = None
        stack: List[Section] = []
        buffer_text: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_level = self._get_style_level(para)
            is_heading, detected_level = self._detect_heading(text, para)

            level = style_level if style_level else detected_level

            if is_heading or style_level:
                # Flush buffer into previous section
                if buffer_text and stack:
                    stack[-1].content += "\n".join(buffer_text)
                buffer_text = []

                new_section = Section(
                    heading=text,
                    level=level if level else 1,
                    content=""
                )

                if not stack:
                    sections.append(new_section)
                    stack = [new_section]
                elif level and level > stack[-1].level:
                    stack[-1].subsections.append(new_section)
                    stack.append(new_section)
                else:
                    while stack and (not level or level <= stack[-1].level):
                        stack.pop()
                    if stack:
                        stack[-1].subsections.append(new_section)
                    else:
                        sections.append(new_section)
                    stack.append(new_section)
            else:
                buffer_text.append(text)

            full_text_parts.append(text)

        # Flush remaining buffer
        if buffer_text and stack:
            stack[-1].content += "\n".join(buffer_text)

        contract = Contract(
            filename=filename,
            file_type="docx",
            full_text="\n".join(full_text_parts),
            sections=sections,
            page_count=paragraph_count // 30 + 1,  # Rough estimate
            metadata=metadata
        )

        return contract

    def _get_style_level(self, para) -> int:
        """Determine heading level from paragraph style."""
        style = para.style
        if style is None:
            return 0
        style_name = style.name.lower() if style.name else ""

        if "heading 1" in style_name:
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        elif "title" in style_name:
            return 1
        elif "subtitle" in style_name:
            return 2
        elif "heading" in style_name:
            return 1
        return 0

    def _detect_heading(self, text: str, para) -> tuple[bool, int]:
        """Detect if text is a heading by pattern matching."""
        if len(text) > 200 or len(text) < 2:
            return False, 0

        # Check font size/bold from runs
        if para.runs:
            first_run = para.runs[0]
            is_bold = first_run.bold
            font_size = first_run.font.size
            if is_bold and font_size and font_size.pt >= 12:
                return True, 1

        # Check for numbered sections
        if re.match(r'^\d+(?:\.\d+)*\.?\s+\w', text):
            num_parts = text.split()[0].rstrip('.')
            dot_count = num_parts.count('.')
            return True, max(1, dot_count + 1)

        # Check for ALL CAPS
        if text.isupper() and len(text) < 100 and len(text.split()) <= 12:
            return True, 1

        # Check for "Section X" pattern
        if re.match(r'^Section\s+\d+', text, re.IGNORECASE):
            return True, 1

        # Check for typical heading keywords
        heading_keywords = [
            'agreement', 'definitions', 'scope', 'terms', 'payment',
            'delivery', 'warranty', 'indemnification', 'confidentiality',
            'termination', 'governing', 'limitation', 'assignment',
            'dispute', 'force majeure', 'insurance', 'representations',
            'covenants', 'general provisions', 'miscellaneous'
        ]

        text_lower = text.lower().rstrip(':')
        for keyword in heading_keywords:
            if text_lower == keyword or text_lower.startswith(keyword):
                return True, 1

        return False, 0